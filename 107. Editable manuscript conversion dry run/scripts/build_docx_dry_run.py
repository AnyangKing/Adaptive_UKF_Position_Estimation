from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = Path(__file__).resolve().parents[1]
SRC = ROOT / "105. Abstract and back-matter tightening" / "manuscript_sensors_candidate_tightened.md"
OUT = OUT_DIR / "editable_manuscript_conversion_dry_run.docx"

FIGURES = {
    "Fig. 1": ROOT / "104. Final submission asset package dry run" / "figures" / "fig1_system_concept_polished.png",
    "Fig. 2": ROOT / "104. Final submission asset package dry run" / "figures" / "fig2_frequency_agile_bias.png",
    "Fig. 3": ROOT / "104. Final submission asset package dry run" / "figures" / "fig3_static_600m_paired_rmse.png",
    "Fig. 4": ROOT / "104. Final submission asset package dry run" / "figures" / "fig4_moving_whitening_lag1.png",
    "Fig. 5": ROOT / "104. Final submission asset package dry run" / "figures" / "fig5_quasi_static_speed_boundary.png",
    "Fig. 6": ROOT / "104. Final submission asset package dry run" / "figures" / "fig6_crlb_floor.png",
}


def set_margins(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5), 18, 10),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5), 14, 7),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78), 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.25

    if "DryRunEquation" not in styles:
        eq = styles.add_style("DryRunEquation", 1)
        eq.font.name = "Consolas"
        eq.font.size = Pt(9)
        eq.paragraph_format.space_before = Pt(4)
        eq.paragraph_format.space_after = Pt(6)
        eq.paragraph_format.line_spacing = 1.0

    if "CaptionDryRun" not in styles:
        cap = styles.add_style("CaptionDryRun", 1)
        cap.font.name = "Calibri"
        cap.font.size = Pt(9)
        cap.font.italic = True
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after = Pt(8)


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "")
    text = text.replace("--", "-")
    return text


def add_para(doc, text, style=None):
    if not text.strip():
        return
    p = doc.add_paragraph(style=style)
    p.add_run(clean_inline(text.strip()))


def parse_table(lines):
    rows = []
    for line in lines:
        if re.match(r"^\|\s*-", line):
            continue
        cells = [clean_inline(c.strip()) for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[j] if j < len(row) else ""
            cell.text = text
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8 if cols >= 5 else 9)
                    if i == 0:
                        run.bold = True
    doc.add_paragraph()


def add_figure_if_caption(doc, line):
    match = re.match(r"\*\*(Fig\. [1-6])\.\*\*", line)
    if not match:
        return False
    key = match.group(1)
    fig = FIGURES.get(key)
    if fig and fig.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(fig), width=Inches(6.2))
    cap = doc.add_paragraph(style="CaptionDryRun")
    cap.add_run(clean_inline(line))
    return True


def build():
    doc = Document()
    set_margins(doc)
    configure_styles(doc)

    lines = SRC.read_text(encoding="utf-8").splitlines()

    # Title
    title = lines[0].lstrip("# ").strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("Editable manuscript conversion dry run - not a final submission file")
    nr.italic = True
    nr.font.size = Pt(9)

    in_math = False
    math_buf = []
    table_buf = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            add_table(doc, parse_table(table_buf))
            table_buf = []

    def flush_math():
        nonlocal math_buf
        if math_buf:
            eq_no = flush_math.counter
            p = doc.add_paragraph(style="DryRunEquation")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"Eq. ({eq_no})  " + " ".join(m.strip() for m in math_buf))
            flush_math.counter += 1
            math_buf = []

    flush_math.counter = 1

    for raw in lines[1:]:
        line = raw.rstrip()
        if line.startswith("```math"):
            flush_table()
            in_math = True
            math_buf = []
            continue
        if in_math:
            if line.startswith("```"):
                in_math = False
                flush_math()
            else:
                math_buf.append(line)
            continue
        if line.startswith("|") and line.endswith("|"):
            table_buf.append(line)
            continue
        flush_table()

        if not line.strip() or line.startswith("> "):
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            doc.add_heading(clean_inline(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(clean_inline(line[4:]), level=2)
        elif add_figure_if_caption(doc, line):
            continue
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(clean_inline(line[2:]))
        else:
            add_para(doc, line)

    flush_table()
    flush_math()

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
